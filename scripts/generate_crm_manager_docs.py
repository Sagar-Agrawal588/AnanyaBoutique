from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from PIL import Image as PilImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import (
    Image as PdfImage,
    ListFlowable,
    ListItem,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "crm-manager-handoff.md"
DOCX_OUT = ROOT / "docs" / "CRM-Manager-Handoff.docx"
PDF_OUT = ROOT / "docs" / "CRM-Manager-Handoff.pdf"
SCREENSHOT_DIR = ROOT / "docs" / "screenshots"
SCREENSHOTS = [
    ("crm-dashboard.png", "CRM dashboard overview"),
    ("crm-whatsapp-view.png", "WhatsApp-filtered CRM timeline"),
]


def parse_markdown(text: str) -> list[dict]:
    blocks: list[dict] = []
    paragraph_lines: list[str] = []
    list_items: list[str] = []
    list_mode: str | None = None
    code_lines: list[str] = []
    in_code = False

    def flush_paragraph() -> None:
      nonlocal paragraph_lines
      if paragraph_lines:
          blocks.append({"type": "paragraph", "text": " ".join(paragraph_lines).strip()})
          paragraph_lines = []

    def flush_list() -> None:
      nonlocal list_items, list_mode
      if list_items:
          blocks.append({"type": "list", "mode": list_mode or "ul", "items": list_items[:]})
          list_items = []
          list_mode = None

    def flush_code() -> None:
      nonlocal code_lines
      if code_lines:
          blocks.append({"type": "code", "text": "\n".join(code_lines).rstrip()})
          code_lines = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            flush_list()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(raw_line)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            flush_list()
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": heading_match.group(2).strip(),
                }
            )
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered_match:
            flush_paragraph()
            if list_mode not in {None, "ol"}:
                flush_list()
            list_mode = "ol"
            list_items.append(ordered_match.group(1).strip())
            continue

        unordered_match = re.match(r"^-\s+(.*)$", stripped)
        if unordered_match:
            flush_paragraph()
            if list_mode not in {None, "ul"}:
                flush_list()
            list_mode = "ul"
            list_items.append(unordered_match.group(1).strip())
            continue

        if not stripped:
            flush_paragraph()
            flush_list()
            continue

        paragraph_lines.append(stripped)

    flush_paragraph()
    flush_list()
    flush_code()
    return blocks


def render_docx(blocks: list[dict], output_path: Path) -> None:
    document = Document()
    document.core_properties.title = "CRM Manager Handoff"
    document.core_properties.subject = "bogEcom CRM implementation handoff"

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    for block in blocks:
        block_type = block["type"]

        if block_type == "heading":
            level = min(int(block["level"]), 4)
            document.add_heading(block["text"], level=level)
            continue

        if block_type == "paragraph":
            document.add_paragraph(block["text"])
            continue

        if block_type == "list":
            style_name = "List Number" if block["mode"] == "ol" else "List Bullet"
            for item in block["items"]:
                document.add_paragraph(item, style=style_name)
            continue

        if block_type == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(9)

    available_screenshots = [
        (SCREENSHOT_DIR / file_name, caption)
        for file_name, caption in SCREENSHOTS
        if (SCREENSHOT_DIR / file_name).exists()
    ]

    if available_screenshots:
        document.add_heading("Local Screenshots", level=2)
        for image_path, caption in available_screenshots:
            document.add_paragraph(caption)
            document.add_picture(str(image_path), width=Inches(6.5))

    document.save(output_path)


def render_pdf(blocks: list[dict], output_path: Path) -> None:
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    heading_styles = {
        1: ParagraphStyle(
            "H1",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            textColor=colors.HexColor("#111827"),
            spaceAfter=10,
        ),
        2: ParagraphStyle(
            "H2",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#1f2937"),
            spaceBefore=6,
            spaceAfter=8,
        ),
        3: ParagraphStyle(
            "H3",
            parent=styles["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#334155"),
            spaceBefore=4,
            spaceAfter=6,
        ),
        4: ParagraphStyle(
            "H4",
            parent=styles["Heading4"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#475569"),
            spaceBefore=4,
            spaceAfter=6,
        ),
    }
    code_style = ParagraphStyle(
        "Code",
        parent=body,
        fontName="Courier",
        fontSize=8.5,
        leading=11,
        backColor=colors.HexColor("#F8FAFC"),
        borderPadding=6,
        borderColor=colors.HexColor("#CBD5E1"),
        borderWidth=0.5,
        spaceAfter=10,
    )

    story = []

    for block in blocks:
        block_type = block["type"]

        if block_type == "heading":
            style = heading_styles.get(min(int(block["level"]), 4), heading_styles[4])
            story.append(Paragraph(html.escape(block["text"]), style))
            continue

        if block_type == "paragraph":
            story.append(Paragraph(html.escape(block["text"]), body))
            continue

        if block_type == "list":
            list_items = [
                ListItem(Paragraph(html.escape(item), body), leftIndent=10)
                for item in block["items"]
            ]
            bullet_type = "1" if block["mode"] == "ol" else "bullet"
            story.append(
                ListFlowable(
                    list_items,
                    bulletType=bullet_type,
                    start="1",
                    leftIndent=14,
                    bulletFontName="Helvetica",
                )
            )
            story.append(Spacer(1, 6))
            continue

        if block_type == "code":
            story.append(Preformatted(block["text"], code_style))

    available_screenshots = [
        (SCREENSHOT_DIR / file_name, caption)
        for file_name, caption in SCREENSHOTS
        if (SCREENSHOT_DIR / file_name).exists()
    ]

    if available_screenshots:
        story.append(Paragraph("Local Screenshots", heading_styles[2]))
        max_width = A4[0] - 80

        for image_path, caption in available_screenshots:
            story.append(Paragraph(html.escape(caption), body))
            with PilImage.open(image_path) as image:
                width_px, height_px = image.size

            scale = min(max_width / float(width_px or 1), 1.0)
            pdf_image = PdfImage(str(image_path))
            pdf_image.drawWidth = width_px * scale
            pdf_image.drawHeight = height_px * scale
            story.append(pdf_image)
            story.append(Spacer(1, 10))

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        title="CRM Manager Handoff",
        author="Codex",
        leftMargin=40,
        rightMargin=40,
        topMargin=36,
        bottomMargin=36,
    )
    document.build(story)


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Source markdown not found: {SOURCE}")

    markdown_text = SOURCE.read_text(encoding="utf-8")
    blocks = parse_markdown(markdown_text)

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    render_docx(blocks, DOCX_OUT)
    render_pdf(blocks, PDF_OUT)

    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {PDF_OUT}")


if __name__ == "__main__":
    main()
