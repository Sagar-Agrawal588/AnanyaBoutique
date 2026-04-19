from docx import Document
from pptx import Presentation


doc = Document(r"d:\BogEcom\Piyush_Songara_Mid_Review_Report.docx")
texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
word_count = sum(len(t.split()) for t in texts)
print("REPORT_PARAGRAPHS", len(texts))
print("REPORT_WORDS", word_count)
print("REPORT_TABLES", len(doc.tables))
print("APPENDIX_COMMIT_LINES", sum(1 for t in texts if "|" in t and "frontend/" not in t))
print("APPENDIX_FILE_LINES", sum(1 for t in texts if t.startswith("frontend/")))

prs = Presentation(r"d:\BogEcom\Piyush_Songara_Mid_Review_Presentation.pptx")
print("PPT_SLIDES", len(prs.slides))
