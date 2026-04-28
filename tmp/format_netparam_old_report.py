from __future__ import annotations

import os
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE = Path(r"C:\Users\piyush songara\Downloads\Mid_Review_Report_Netparam_Formatted old.docx")
BACKUP = SOURCE.with_name("Mid_Review_Report_Netparam_Formatted old.backup.docx")
FALLBACK = SOURCE.with_name("Mid_Review_Report_Netparam_Formatted_University_Format.docx")
NETPARAM_LOGO = Path(r"D:\BogEcom\output\friend-report-media\netparam_logo_generated.png")
JECRC_LOGO = Path(r"D:\BogEcom\output\template1-media\image1.jpeg")


MAIN_SECTION_HEADINGS = {
    "DECLARATION",
    "ACKNOWLEDGEMENT",
    "INDEX",
    "LIST OF ILLUSTRATIONS",
    "ABSTRACT",
    "REFERENCES",
}


DECLARATION_TEXT = (
    'I hereby declare that the internship done at Netparam Technologies Pvt. Ltd. '
    'on the project "Perfume Notes Classification System" is an authentic record '
    "of my own work carried out as a requirement of the mid review internship for "
    "the award of the degree of B.Tech in Computer Science and Engineering, "
    "JECRC University, Jaipur, under the guidance of the industry coordinator "
    "and faculty guide mentioned in this report."
)

DECLARATION_TEXT_2 = (
    "The contents of this report are based on the actual implementation work "
    "completed during the internship period, including dataset preparation, "
    "machine learning workflow design, model training, Flask application "
    "integration, testing, and documentation. This report has not been submitted "
    "in the same form for any other academic requirement."
)

ACK_TEXT_1 = (
    "It is my foremost duty to express my deep sense of gratitude and respect "
    "to [Faculty Guide Name], Assistant Professor, for constant encouragement, "
    "guidance, and support in making this internship work successful. I am also "
    "grateful to the Department of Computer Science and Engineering, JECRC "
    "University, Jaipur, for providing the academic environment and necessary "
    "facilities to carry out this internship work effectively."
)

ACK_TEXT_2 = (
    "I am also thankful to [Industry Guide Name], [Industry Guide Designation], "
    "for providing me with the opportunity to complete my internship at "
    "Netparam Technologies Pvt. Ltd. The project work in the ALCHEMIST "
    "environment helped me understand how data preparation, model development, "
    "application integration, and project presentation come together in a "
    "professional workflow."
)

ACK_TEXT_3 = (
    "Last but not least, I am thankful to my friends, colleagues, and everyone "
    "who directly or indirectly supported me during the internship period."
)

LOI_TEXT = (
    "The list of illustrations gives systematic information about tables, figures, "
    "charts, and other visual material used in the report. It provides a "
    "structured overview of the illustrations included to support the "
    "technical explanation presented across different sections."
)


def set_run_font(run, size=12, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size=12, bold=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def replace_text(paragraph, text):
    if not paragraph.runs:
        paragraph.text = text
        set_paragraph_font(paragraph, size=12)
        return
    first = paragraph.runs[0]
    first.text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def is_main_heading(text: str) -> bool:
    clean = text.strip().upper()
    return clean in MAIN_SECTION_HEADINGS or bool(re.match(r"^\d+\s+[A-Z]", clean))


def is_sub_heading(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\s+", text.strip()))


def is_cover_paragraph(index: int) -> bool:
    return index <= 18


def is_signature_line(text: str) -> bool:
    t = text.strip()
    return (
        t.startswith("Name of Student:")
        or t.startswith("RN No.")
        or t.startswith("Place:")
        or t.startswith("Date:")
        or t.endswith("Guide")
        or t in {"(Sign of Student)", "Thanking You.", "[Your Name]"}
    )


def format_cover_paragraph(paragraph, idx):
    text = paragraph.text.strip()
    fmt = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Inches(0)
    size = 12
    bold = False
    if text == "MID REVIEW INTERNSHIP REPORT":
        size = 16
        bold = True
    elif text == "(Internship Semester Jan-June 2026)":
        size = 13
        bold = True
    elif text == "BACHELOR OF TECHNOLOGY":
        size = 14
        bold = True
    elif text == "COMPUTER SCIENCE AND ENGINEERING":
        size = 13
        bold = True
    elif text in {"By", "Under the Guidance of"}:
        size = 12
        bold = True
    elif text in {"[Your Name]", "Department of Computer Science and Engineering", "JECRC UNIVERSITY, JAIPUR"}:
        size = 12
        bold = True
    elif text.startswith("Reg. No."):
        size = 12
        bold = True
    set_paragraph_font(paragraph, size=size, bold=bold)


def format_paragraph(paragraph, idx):
    text = paragraph.text.strip()
    fmt = paragraph.paragraph_format
    fmt.left_indent = Inches(0)
    fmt.right_indent = Inches(0)
    fmt.first_line_indent = Inches(0)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)

    if not text:
        return

    if is_cover_paragraph(idx):
        format_cover_paragraph(paragraph, idx)
        return

    if text == "DECLARATION":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_font(paragraph, size=16, bold=True)
        return

    if text == "ACKNOWLEDGEMENT":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_font(paragraph, size=16, bold=True)
        return

    if text in {"INDEX", "LIST OF ILLUSTRATIONS", "ABSTRACT", "REFERENCES"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_font(paragraph, size=16, bold=True)
        return

    if re.match(r"^\d+\s+[A-Z]", text):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_font(paragraph, size=16, bold=True)
        return

    if is_sub_heading(text):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_font(paragraph, size=14, bold=True)
        return

    if text.startswith("Figure ") or text.startswith("Table "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_font(paragraph, size=11, bold=True)
        return

    if is_signature_line(text):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_font(paragraph, size=12)
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_font(paragraph, size=12)


def format_tables(doc):
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(3)
                    para.paragraph_format.space_before = Pt(0)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    set_paragraph_font(para, size=12, bold=(row_idx == 0))


def ensure_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.left_margin = Inches(1.25)


def normalize_styles(doc):
    for style_name in ["Normal", "Body Text"]:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)


def apply_front_matter_text_updates(doc):
    mapping = {
        20: "DECLARATION",
        21: DECLARATION_TEXT,
        22: DECLARATION_TEXT_2,
        30: "ACKNOWLEDGEMENT",
        31: ACK_TEXT_1,
        32: ACK_TEXT_2,
        33: ACK_TEXT_3,
        39: "List of Illustrations",
        40: LOI_TEXT,
    }
    for idx, text in mapping.items():
        if idx < len(doc.paragraphs):
            replace_text(doc.paragraphs[idx], text)


def cleanup_front_matter_duplicates(doc):
    normalized = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            normalized.append((p, t))

    for i in range(1, len(normalized)):
        prev_p, prev_t = normalized[i - 1]
        p, t = normalized[i]
        if prev_t == "DECLARATION" and t == "DECLARATION":
            replace_text(p, "")
        if prev_t == "ACKNOWLEDGEMENT" and t == "ACKNOWLEDGEMENT":
            replace_text(p, "")
        if prev_t == "LIST OF ILLUSTRATIONS" and t == "List of Illustrations":
            replace_text(prev_p, "")


def replace_media(docx_path: Path, replacements: dict[str, Path]):
    temp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            replacement = replacements.get(item.filename)
            if replacement and replacement.exists():
                data = replacement.read_bytes()
            zout.writestr(item, data)
    temp.replace(docx_path)


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    shutil.copy2(SOURCE, BACKUP)

    doc = Document(str(SOURCE))
    normalize_styles(doc)
    ensure_margins(doc)
    apply_front_matter_text_updates(doc)
    cleanup_front_matter_duplicates(doc)

    for idx, paragraph in enumerate(doc.paragraphs):
        format_paragraph(paragraph, idx)

    format_tables(doc)
    output = SOURCE
    try:
        doc.save(str(output))
    except PermissionError:
        output = FALLBACK
        doc.save(str(output))
    replace_media(
        output,
        {
            "word/media/image1.jpeg": JECRC_LOGO,
            "word/media/image2.png": NETPARAM_LOGO,
        },
    )
    print(f"Formatted: {output}")
    print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
