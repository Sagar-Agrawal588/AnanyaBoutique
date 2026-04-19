from __future__ import annotations

import subprocess
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


ROOT = Path(r"d:\BogEcom")
OUTPUT = ROOT / "Piyush_Work_Summary_2026-03-08_to_2026-04-08.docx"


def git(*args: str) -> str:
    result = subprocess.run(
        ["git", *args],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    return result.stdout


def build_data() -> dict[str, object]:
    base_args = [
        "--since=2026-03-08",
        "--until=2026-04-08",
        "--author=piyushsongara69@gmail.com",
        "--",
        "frontend/client",
        "frontend/admin",
    ]

    commits = [
        line.strip()
        for line in git("log", "--oneline", *base_args).splitlines()
        if line.strip()
    ]
    count = int(git("rev-list", "--count", "--since=2026-03-08", "--until=2026-04-08", "--author=piyushsongara69@gmail.com", "HEAD", "--", "frontend/client", "frontend/admin").strip())
    files = [
        line.strip()
        for line in git("log", "--name-only", "--pretty=tformat:", *base_args).splitlines()
        if line.strip()
    ]
    file_counter = Counter(files)

    additions = 0
    deletions = 0
    for line in git("log", "--numstat", "--pretty=tformat:", *base_args).splitlines():
        parts = line.split("\t")
        if len(parts) != 3:
            continue
        add, delete, _ = parts
        if add != "-":
            additions += int(add)
        if delete != "-":
            deletions += int(delete)

    categories = Counter()
    for file_path in sorted(set(files)):
        normalized = file_path.replace("\\", "/")
        if "/src/app/" in normalized:
            categories["Pages and flows"] += 1
        elif "/src/components/" in normalized:
            categories["UI components"] += 1
        elif "/src/context/" in normalized:
            categories["State management"] += 1
        elif "/src/utils/" in normalized:
            categories["Utilities and API"] += 1
        elif "/src/hooks/" in normalized:
            categories["Hooks and integration"] += 1
        else:
            categories["Config and support files"] += 1

    top_files = file_counter.most_common(6)
    recent = commits[:8]

    return {
        "commit_count": count,
        "unique_files": len(set(files)),
        "additions": additions,
        "deletions": deletions,
        "top_files": top_files,
        "recent_commits": recent,
        "categories": categories,
    }


def set_normal(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_para(doc: Document, text: str, *, bold: bool = False, size: float = 10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def build_doc(data: dict[str, object]) -> None:
    doc = Document()
    set_normal(doc)

    add_para(doc, "ONE-PAGE WORK SUMMARY", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_para(doc, "Piyush Songara | Project Contribution from March 8, 2026 to April 8, 2026", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    add_para(
        doc,
        f"Between March 8, 2026 and April 8, 2026, I contributed actively to the BogEcom project mainly in frontend/client and frontend/admin. Git history for this exact period shows {data['commit_count']} authored frontend-related commits, touching {data['unique_files']} unique files, with approximately {data['additions']} lines added and {data['deletions']} lines deleted. This reflects sustained work on production-facing ecommerce flows rather than isolated one-time changes.",
    )

    add_para(
        doc,
        "My contribution during this period was concentrated on customer-facing flows such as product detail pages, checkout, cart, my orders, membership, product cards, sliders, and shared UI components, along with selected admin-side modules such as orders, settings, shipping, purchase orders, combos, sidebar, and analytics-related interfaces.",
    )

    add_para(doc, "Major work areas in this period:", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)
    add_bullet(doc, "Improved product presentation through updates to ProductItem, ProductSlider, PopularProducts, Product Detail, image handling, and responsive layout behavior.")
    add_bullet(doc, "Fixed order and checkout issues including order-id flow, cart totals, combo pricing, PhonePe checkout behavior, and related frontend calculation alignment.")
    add_bullet(doc, "Worked on reusable app behavior through Header, CartDrawer, Wishlist, Settings, CartContext, and supporting utility/API modules.")
    add_bullet(doc, "Contributed to admin-facing screens such as orders, settings, shipping, purchase orders, combos, and supporting dashboard components.")

    category_text = ", ".join(f"{label}: {count}" for label, count in data["categories"].most_common())
    add_para(doc, f"Contribution spread by file area: {category_text}.", space_after=2)

    add_para(doc, "Most frequently touched files in this window:", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)
    for file_path, count in data["top_files"]:
        add_bullet(doc, f"{file_path} ({count} touches)")

    add_para(doc, "Representative commits from this period:", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)
    for line in data["recent_commits"]:
        add_bullet(doc, line)

    add_para(
        doc,
        "Overall, this one-month period shows that my role was strongly frontend-oriented and involved both UI refinement and logic-sensitive ecommerce work. The work covered high-impact user journeys, shared components, and operational admin screens, demonstrating meaningful and continuous contribution to the project during March 8, 2026 to April 8, 2026.",
        space_after=0,
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc(build_data())
    print(OUTPUT)
