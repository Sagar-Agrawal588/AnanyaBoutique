from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\Users\piyush songara\Downloads\Piyush_Songara_Mid_Review_Report 69.docx")
TARGET = Path(r"C:\Users\piyush songara\Downloads\Mid_Review_Report_Netparam_Formatted_University_Format.docx")
FALLBACK = TARGET.with_name("Mid_Review_Report_Netparam_PiyushStyle.docx")


def copy_run_style(src_run, dst_run):
    dst_run.font.name = src_run.font.name or "Times New Roman"
    dst_run._element.rPr.rFonts.set(qn("w:eastAsia"), src_run.font.name or "Times New Roman")
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline


def paragraph_sample_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def copy_paragraph_format(src, dst):
    dst.style = src.style
    dst.alignment = src.alignment
    spf = src.paragraph_format
    dpf = dst.paragraph_format
    dpf.left_indent = spf.left_indent
    dpf.right_indent = spf.right_indent
    dpf.first_line_indent = spf.first_line_indent
    dpf.space_before = spf.space_before
    dpf.space_after = spf.space_after
    dpf.line_spacing = spf.line_spacing
    dpf.keep_together = spf.keep_together
    dpf.keep_with_next = spf.keep_with_next
    dpf.page_break_before = spf.page_break_before
    dpf.widow_control = spf.widow_control
    sample = paragraph_sample_run(src)
    if sample:
        for run in dst.runs:
            copy_run_style(sample, run)


def copy_table_style(src_table, dst_table):
    src_tbl = src_table._tbl
    dst_tbl = dst_table._tbl

    src_tblPr = src_tbl.tblPr
    if dst_tbl.tblPr is not None:
        dst_tbl.remove(dst_tbl.tblPr)
    dst_tbl.insert(0, deepcopy(src_tblPr))

    if src_tbl.tblGrid is not None:
        if dst_tbl.tblGrid is not None:
            dst_tbl.remove(dst_tbl.tblGrid)
        dst_tbl.insert(1, deepcopy(src_tbl.tblGrid))

    for r_idx, row in enumerate(dst_table.rows):
        for c_idx, cell in enumerate(row.cells):
            if r_idx >= len(src_table.rows) or c_idx >= len(src_table.rows[r_idx].cells):
                continue
            src_cell = src_table.rows[r_idx].cells[c_idx]
            for p_idx, para in enumerate(cell.paragraphs):
                if p_idx < len(src_cell.paragraphs):
                    copy_paragraph_format(src_cell.paragraphs[p_idx], para)


def find_paragraph(doc, text):
    want = re.sub(r"\s+", " ", text.strip())
    for p in doc.paragraphs:
        have = re.sub(r"\s+", " ", p.text.strip())
        if have == want:
            return p
    raise ValueError(f"Paragraph not found: {text}")


def find_paragraph_contains(doc, text):
    want = re.sub(r"\s+", " ", text.strip())
    for p in doc.paragraphs:
        have = re.sub(r"\s+", " ", p.text.strip())
        if want in have:
            return p
    raise ValueError(f"Paragraph containing not found: {text}")


def style_cover(source, target):
    src_cover_keys = [
        "MID REVIEW INTERNSHIP REPORT",
        "(Internship Semester Jan-June 2026)",
        "A report submitted in partial fulfillment of the requirements for the Award of Degree of",
        "BACHELOR OF TECHNOLOGY",
        "in",
        "COMPUTER SCIENCE AND ENGINEERING",
        "By",
        "Reg. No.: 22BCON386",
        "Under the Guidance of",
        "Department of Computer Science and Engineering",
        "JECRC UNIVERSITY, JAIPUR",
        "April 2026",
    ]

    target_map = {
        "MID REVIEW INTERNSHIP REPORT": "MID REVIEW INTERNSHIP REPORT",
        "(Internship Semester Jan-June 2026)": "(Internship Semester Jan-June 2026)",
        "A report submitted in partial fulfillment of the requirements for the Award of Degree of": "A report submitted in partial fulfillment of the requirements for the Award of Degree of",
        "BACHELOR OF TECHNOLOGY": "BACHELOR OF TECHNOLOGY",
        "in": "in",
        "COMPUTER SCIENCE AND ENGINEERING": "COMPUTER SCIENCE AND ENGINEERING",
        "By": "By",
        "Reg. No.: 22BCON386": "Reg. No.: [Your Registration Number]",
        "Under the Guidance of": "Under the Guidance of",
        "Department of Computer Science and Engineering": "Department of Computer Science and Engineering",
        "JECRC UNIVERSITY, JAIPUR": "JECRC UNIVERSITY, JAIPUR",
        "April 2026": "April 2026",
    }

    for src_text in src_cover_keys:
        src_p = find_paragraph(source, src_text)
        dst_p = find_paragraph(target, target_map[src_text])
        copy_paragraph_format(src_p, dst_p)

    # Apply name/guidance row formatting using semantically similar lines.
    copy_paragraph_format(find_paragraph(source, "Piyush Songara"), find_paragraph(target, "[Your Name]"))
    copy_paragraph_format(find_paragraph(source, "Faculty Internship Guide Industry Guide"), find_paragraph(target, "Faculty Internship Guide Industry Guide"))
    copy_paragraph_format(find_paragraph(source, "Name : Ms. Megha Garg Name: Mr. Balveer Bajiya"), find_paragraph(target, "Name: [Faculty Guide Name] Name: [Industry Guide Name]"))
    copy_paragraph_format(find_paragraph(source, "Designation: Assistant Professor Designation: CTO"), find_paragraph(target, "Designation: Assistant Professor Designation: [Industry Guide Designation]"))


def style_section_by_samples(source, target, section_heading, source_samples, target_texts):
    heading_src = find_paragraph(source, section_heading)
    heading_dst = find_paragraph(target, section_heading if section_heading != "LIST OF ILLUSTRATIONS" else "List of Illustrations")
    copy_paragraph_format(heading_src, heading_dst)
    for src_text, dst_text in zip(source_samples, target_texts):
        copy_paragraph_format(find_paragraph(source, src_text), find_paragraph(target, dst_text))


def main():
    src = Document(str(SOURCE))
    dst = Document(str(TARGET))

    style_cover(src, dst)

    style_section_by_samples(
        src,
        dst,
        "DECLARATION",
        [
            'I hereby declare that the internship work carried out on the project "AnanyaBoutique / Ananya Boutique E-commerce Platform" is an authentic record of my work completed as part of the six-month internship requirement for the award of B.Tech in Computer Science and Engineering. The work was performed under the guidance of Balveer Bajiya(CTO, Ananya Boutique) and Megha Garg(Assistant Professor, JECRC University).',
            "Name of Student: Piyush Songara",
            "RN No.: 22BCON386",
            "Place: JECRC University, Jaipur",
            "Date: 23 April 2026",
            "Mr.Balveer Bajiya\t\tGuide",
        ],
        [
            'I hereby declare that the internship done at Netparam Technologies Pvt. Ltd. on the project "Perfume Notes Classification System" is an authentic record of my own work carried out as a requirement of the mid review internship for the award of the degree of B.Tech in Computer Science and Engineering, JECRC University, Jaipur, under the guidance of the industry coordinator and faculty guide mentioned in this report.',
            "The contents of this report are based on the actual implementation work completed during the internship period, including dataset preparation, machine learning workflow design, model training, Flask application integration, testing, and documentation. This report has not been submitted in the same form for any other academic requirement.",
            "Name of Student: [Your Name]",
            "RN No.: [Your Registration Number]",
            "Place: JECRC University, Jaipur",
            "Date: 23 April 2026",
            "[Industry Guide Name]    Guide",
        ],
    )

    style_section_by_samples(
        src,
        dst,
        "ACKNOWLEDGEMENT",
        [
            "It is my foremost duty to express my deep sense of gratitude and respect to the guide Ms. Megha Garg (Assistant Professor) his uplifting tendency and inspiring me for taking up this Internship completely successful. I am also grateful to Dr. Bhavana Sharma (H.O.D), Dr. Naveen Hemrajani (Dean Engineering) and Dy. HoD Dr. Gajanand Sharma for providing all necessary facilities to carry out the internship work and those encouraging part has been a perpetual source of inspiration.",
            "I am also thankful to Mr. Balveer Bajiya (CTO, Ananya Boutique) Sir for proving me internship in Ananya Boutique, Sitapura, Jaipur. Last but not least I am thankful to my colleagues and those help me directly or indirectly through this Internship.",
            "Thanking You.",
            "Piyush Songara",
        ],
        [
            "It is my foremost duty to express my deep sense of gratitude and respect to [Faculty Guide Name], Assistant Professor, for constant encouragement, guidance, and support in making this internship work successful. I am also grateful to the Department of Computer Science and Engineering, JECRC University, Jaipur, for providing the academic environment and necessary facilities to carry out this internship work effectively.",
            "I am also thankful to [Industry Guide Name], [Industry Guide Designation], for providing me with the opportunity to complete my internship at Netparam Technologies Pvt. Ltd. The project work in the ALCHEMIST environment helped me understand how data preparation, model development, application integration, and project presentation come together in a professional workflow.",
            "Last but not least, I am thankful to my friends, colleagues, and everyone who directly or indirectly supported me during the internship period.",
            "[Your Name]",
        ],
    )

    copy_paragraph_format(find_paragraph(src, "INDEX"), find_paragraph(dst, "INDEX"))
    copy_paragraph_format(find_paragraph(src, "LIST OF ILLUSTRATIONS"), find_paragraph(dst, "List of Illustrations"))
    copy_paragraph_format(
        find_paragraph(src, "The list of illustrations gives systematic information about the tables and figures used in this report. It provides a structured overview of where important visual summaries and chapter support material appear in the document."),
        find_paragraph_contains(dst, "The list of illustrations gives systematic information"),
    )

    copy_table_style(src.tables[0], dst.tables[0])
    copy_table_style(src.tables[1], dst.tables[1])

    try:
        dst.save(str(TARGET))
        print(f"Updated: {TARGET}")
    except PermissionError:
        dst.save(str(FALLBACK))
        print(f"Updated: {FALLBACK}")


if __name__ == "__main__":
    main()
